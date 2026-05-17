from dataclasses import dataclass
from pathlib import Path

from fastapi import HTTPException, status


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int | None
    text: str


def validate_file_extension(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type. Upload PDF, TXT, or DOCX files only.",
        )
    return suffix.lstrip(".")


def extract_text(file_path: str, file_type: str) -> list[ExtractedPage]:
    if file_type == "pdf":
        return _extract_pdf(file_path)
    if file_type == "txt":
        return _extract_txt(file_path)
    if file_type == "docx":
        return _extract_docx(file_path)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type")


def _extract_pdf(file_path: str) -> list[ExtractedPage]:
    try:
        import fitz
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="PyMuPDF is not installed") from exc

    pages: list[ExtractedPage] = []
    with fitz.open(file_path) as document:
        for index, page in enumerate(document, start=1):
            pages.append(ExtractedPage(page_number=index, text=page.get_text("text").strip()))
    return pages


def _extract_txt(file_path: str) -> list[ExtractedPage]:
    raw = Path(file_path).read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")
    return [ExtractedPage(page_number=None, text=text.strip())]


def _extract_docx(file_path: str) -> list[ExtractedPage]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is not installed") from exc

    document = DocxDocument(file_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                paragraphs.append(" | ".join(values))
    return [ExtractedPage(page_number=None, text="\n".join(paragraphs).strip())]

